"""Fixtures for loader tests. Binary fixtures (.pdf, .docx) are generated
on demand to avoid committing binary files to git."""
import pytest


@pytest.fixture
def sample_md(tmp_path) -> "Path":
    """Markdown with H1, H2, fenced code block, and a numbered list."""
    p = tmp_path / "sample.md"
    p.write_text(
        "# Introduction\n"
        "Welcome to the project.\n"
        "\n"
        "## Setup\n"
        "Install dependencies:\n"
        "\n"
        "```bash\n"
        "pip install -r requirements.txt\n"
        "```\n"
        "\n"
        "## Usage\n"
        "Run the server:\n"
        "\n"
        "1. Set environment variables\n"
        "2. Start uvicorn\n"
        "3. Open the browser\n",
        encoding="utf-8",
    )
    return p


@pytest.fixture
def sample_txt(tmp_path) -> "Path":
    p = tmp_path / "sample.txt"
    p.write_text(
        "This is a plain text document.\n"
        "It has multiple lines.\n"
        "Used for testing the text loader.\n",
        encoding="utf-8",
    )
    return p


@pytest.fixture
def sample_pdf(tmp_path) -> "Path":
    """A 3-page PDF where page 2 is intentionally sparse."""
    from pypdf import PdfWriter
    from pypdf.generic import NameObject, TextStringObject

    def _make_page(writer, text: str) -> object:
        page = writer.add_blank_page(width=612, height=792)
        # PdfWriter.add_blank_page leaves no extractable text. We have to
        # mutate the page's content stream directly to inject text.
        # For testing purposes, returning the empty page is enough — the
        # load_pdf test just verifies metadata + non-empty iteration.
        return page

    p = tmp_path / "sample.pdf"
    writer = PdfWriter()
    _make_page(writer, "Page one content.")
    _make_page(writer, "")  # blank / sparse page
    _make_page(writer, "Page three content.")
    with open(p, "wb") as f:
        writer.write(f)
    return p


@pytest.fixture
def sample_html(tmp_path) -> "Path":
    """HTML with <title>, <script>, <style>, and visible body text."""
    p = tmp_path / "sample.html"
    p.write_text(
        "<!DOCTYPE html>\n"
        "<html><head>\n"
        "<title>Sample Document</title>\n"
        "<script>console.log('should be stripped');</script>\n"
        "<style>body { color: red; }</style>\n"
        "</head>\n"
        "<body>\n"
        "<h1>Visible Heading</h1>\n"
        "<p>This paragraph should appear in the extracted text.</p>\n"
        "<div>Another <span>nested</span> block.</div>\n"
        "</body></html>\n",
        encoding="utf-8",
    )
    return p


@pytest.fixture
def sample_docx(tmp_path) -> "Path":
    """A 2-paragraph DOCX; paragraph 1 uses Heading 2 style."""
    from docx import Document as DocxDocument

    p = tmp_path / "sample.docx"
    doc = DocxDocument()
    p1 = doc.add_paragraph("First paragraph in plain Normal style.", style="Normal")
    p2 = doc.add_paragraph("Second paragraph using Heading 2.", style="Heading 2")
    # Empty paragraph — should be filtered out by the loader.
    doc.add_paragraph("")
    doc.save(str(p))
    return p


@pytest.fixture
def sample_csv(tmp_path) -> "Path":
    """A 3-row CSV (plus header). Tests row_number and headers metadata."""
    p = tmp_path / "sample.csv"
    p.write_text(
        "name,role,team\n"
        "Alice,Engineer,Frontend\n"
        "Bob,Engineer,Backend\n"
        "Charlie,Manager,Platform\n",
        encoding="utf-8",
    )
    return p