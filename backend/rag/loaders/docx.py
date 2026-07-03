"""Loader for .docx files. Emits one RawDocument per non-empty paragraph
with paragraph_number and style metadata. Tables and embedded objects
are intentionally skipped — they're out of scope for v1.
"""
from pathlib import Path
from typing import Iterator

from backend.rag.loaders import RawDocument, register


@register(".docx")
def load_docx(path: Path, source: str) -> Iterator[RawDocument]:
    from docx import Document as DocxDocument
    doc = DocxDocument(str(path))
    for i, para in enumerate(doc.paragraphs, start=1):
        text = para.text
        if not text.strip():
            continue
        style = para.style.name if para.style else ""
        yield RawDocument(
            text=text,
            metadata={
                "format": ".docx",
                "paragraph_number": i,
                "style": style,
            },
        )