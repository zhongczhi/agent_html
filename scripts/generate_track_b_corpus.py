"""Generate a small heterogeneous-format RAG corpus for Track B.

The corpus mimics a real industrial scenario: a fictional company
("TechCorp Inc.") with documents in 6 different file formats. Each
file is a real file of the named format (PDF, DOCX, HTML, CSV, MD,
TXT) — the loaders in `backend.rag.loaders` are exercised end-to-end,
not just text extraction.

Output directory: `scripts/.cache/track_b_corpus/`
  - annual_report.pdf        (3 pages with a table)
  - hr_handbook.docx         (Word doc with headings)
  - employees.csv            (employee list, 20 rows)
  - pricing.md               (markdown with a table)
  - faq.md                   (markdown FAQ)
  - terms.html               (HTML ToS)
  - README.txt               (plain text)

The QA pairs in `track_b_qa_pairs.json` (next to this script) reference
specific facts in these files. The companion eval script
`scripts/eval_track_b.py` runs the SOTA pipeline on the corpus and
reports per-format `contains_gold`.
"""
from __future__ import annotations

import csv
import json
import random
from pathlib import Path

# Output directory
SCRIPT_DIR = Path(__file__).resolve().parent
CORPUS_DIR = SCRIPT_DIR / ".cache" / "track_b_corpus"
CORPUS_DIR.mkdir(parents=True, exist_ok=True)

# Fixed seed for reproducibility
random.seed(42)


# ── 1. annual_report.pdf ────────────────────────────────────────────────────
def make_annual_report_pdf(out: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
    )

    doc = SimpleDocTemplate(
        str(out),
        pagesize=letter,
        title="TechCorp Inc. 2025 Annual Report",
        author="TechCorp Inc.",
    )
    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    body = styles["BodyText"]

    flow = []
    # Page 1: cover + summary
    flow.append(Paragraph("TechCorp Inc. 2025 Annual Report", h1))
    flow.append(Spacer(1, 0.3 * inch))
    flow.append(Paragraph(
        "Fiscal year 2025 marked a year of disciplined growth for TechCorp. "
        "Total revenue reached <b>$48.7 million</b>, up 23% year-over-year. "
        "The customer base grew to <b>1,247 active enterprise accounts</b> "
        "across 38 countries.",
        body,
    ))
    flow.append(Spacer(1, 0.2 * inch))
    flow.append(Paragraph(
        "We ended the year with 312 full-time employees, of which 184 were "
        "engineers and 41 were in marketing. Headcount grew 18% over 2024.",
        body,
    ))
    flow.append(PageBreak())

    # Page 2: financial table
    flow.append(Paragraph("Financial Summary (USD millions)", h2))
    flow.append(Spacer(1, 0.2 * inch))
    data = [
        ["Metric", "FY 2024", "FY 2025", "YoY Change"],
        ["Revenue", "39.6", "48.7", "+22.9%"],
        ["Gross Profit", "27.1", "34.2", "+26.2%"],
        ["Operating Income", "4.8", "7.1", "+47.9%"],
        ["Net Income", "3.2", "4.9", "+53.1%"],
        ["Cash & Equivalents", "21.4", "26.8", "+25.2%"],
    ]
    table = Table(data, colWidths=[2 * inch, 1.2 * inch, 1.2 * inch, 1.2 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    flow.append(table)
    flow.append(Spacer(1, 0.3 * inch))
    flow.append(Paragraph(
        "Looking forward, our 2026 revenue target is <b>$62 million</b>, with a "
        "planned expansion of the engineering team by 60 net hires.",
        body,
    ))
    flow.append(PageBreak())

    # Page 3: leadership
    flow.append(Paragraph("Leadership", h2))
    flow.append(Spacer(1, 0.2 * inch))
    flow.append(Paragraph(
        "<b>Jane Doe</b> — Chief Executive Officer. Jane co-founded TechCorp in "
        "2018 and has led the company through three funding rounds totaling "
        "$71 million.",
        body,
    ))
    flow.append(Spacer(1, 0.1 * inch))
    flow.append(Paragraph(
        "<b>Marcus Chen</b> — Chief Technology Officer. Marcus leads engineering, "
        "data, and infrastructure. He joined as employee #7 in 2019.",
        body,
    ))
    flow.append(Spacer(1, 0.1 * inch))
    flow.append(Paragraph(
        "<b>Priya Patel</b> — Chief Financial Officer. Priya joined in 2022 "
        "from a Big Four accounting firm.",
        body,
    ))
    doc.build(flow)


# ── 2. hr_handbook.docx ────────────────────────────────────────────────────
def make_hr_handbook_docx(out: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.core_properties.title = "TechCorp Employee Handbook 2025"
    doc.core_properties.author = "TechCorp HR"

    doc.add_heading("TechCorp Employee Handbook 2025", level=1)
    doc.add_paragraph(
        "This handbook covers the policies and benefits that apply to all "
        "full-time TechCorp employees as of January 2025."
    )

    doc.add_heading("1. Paid Time Off", level=2)
    doc.add_paragraph(
        "All full-time employees receive 20 days of paid vacation per year, "
        "accruing at 1.67 days per month. Unused vacation carries over up to "
        "a maximum of 5 days into the next calendar year."
    )
    doc.add_paragraph(
        "Senior engineers (level L5 and above) receive an additional 5 days, "
        "for a total of 25 days per year."
    )

    doc.add_heading("2. Parental Leave", level=2)
    doc.add_paragraph(
        "Primary caregivers receive 16 weeks of fully paid parental leave. "
        "Secondary caregivers receive 8 weeks. Both can be taken any time "
        "within the first 12 months after birth or adoption."
    )

    doc.add_heading("3. Remote Work", level=2)
    doc.add_paragraph(
        "TechCorp operates as a hybrid workplace. Employees are expected to "
        "be in the office on Tuesdays and Wednesdays. The other three days "
        "are flexible. Fully remote arrangements are reviewed case-by-case "
        "and require VP-level approval."
    )

    doc.add_heading("4. Expense Reimbursement", level=2)
    doc.add_paragraph(
        "Business expenses under $50 do not require pre-approval. Expenses "
        "between $50 and $500 require manager approval. Expenses over $500 "
        "require finance team approval. Submit reimbursements within 30 days."
    )

    doc.add_heading("5. Code of Conduct", level=2)
    doc.add_paragraph(
        "TechCorp is committed to a workplace free from harassment and "
        "discrimination. Report any concerns to your manager or to HR "
        "directly at conduct@techcorp.example. All reports are confidential."
    )
    doc.save(str(out))


# ── 3. employees.csv ──────────────────────────────────────────────────────
def make_employees_csv(out: Path) -> None:
    rows = [
        ["name", "role", "department", "level", "location", "start_year"],
        ["Jane Doe", "CEO", "Executive", "M3", "San Francisco", 2018],
        ["Marcus Chen", "CTO", "Engineering", "M2", "San Francisco", 2019],
        ["Priya Patel", "CFO", "Finance", "M2", "New York", 2022],
        ["Sofia Rodriguez", "VP Engineering", "Engineering", "M1", "Remote", 2020],
        ["David Kim", "VP Marketing", "Marketing", "M1", "New York", 2021],
        ["Aisha Khan", "VP Sales", "Sales", "M1", "San Francisco", 2020],
        ["Tom Wilson", "Senior Engineer", "Engineering", "L6", "Remote", 2020],
        ["Yuki Tanaka", "Senior Engineer", "Engineering", "L6", "Tokyo", 2021],
        ["Carlos Mendoza", "Senior Engineer", "Engineering", "L5", "Remote", 2022],
        ["Hannah Schmidt", "Engineer", "Engineering", "L4", "Berlin", 2023],
        ["Liam O'Brien", "Engineer", "Engineering", "L4", "Dublin", 2023],
        ["Zara Ahmed", "Engineer", "Engineering", "L3", "San Francisco", 2024],
        ["Mateo Rossi", "Engineer", "Engineering", "L3", "Remote", 2024],
        ["Ananya Iyer", "Marketing Manager", "Marketing", "L5", "New York", 2022],
        ["Chen Wei", "Marketing Manager", "Marketing", "L5", "San Francisco", 2023],
        ["Elena Petrova", "Sales Director", "Sales", "L5", "London", 2021],
        ["Jamal Brown", "Account Executive", "Sales", "L4", "New York", 2023],
        ["Olivia Garcia", "Account Executive", "Sales", "L3", "Remote", 2024],
        ["Ravi Krishnan", "Data Scientist", "Data", "L5", "Remote", 2022],
        ["Linnea Berg", "Data Analyst", "Data", "L3", "Stockholm", 2024],
    ]
    with out.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(rows)


# ── 4. pricing.md ──────────────────────────────────────────────────────────
def make_pricing_md(out: Path) -> None:
    out.write_text(
        """# TechCorp Pricing

Last updated: January 2025.

## Plans

| Plan       | Monthly price | Seats included | API requests / month | Support SLA |
|------------|---------------|----------------|----------------------|-------------|
| Free       | $0            | 1              | 1,000                | Community   |
| Team       | $49           | 5              | 50,000               | 48h email   |
| Business   | $499          | 25             | 500,000              | 24h email   |
| Enterprise | Custom        | Unlimited      | Custom               | 1h phone    |

## Add-ons

- **Additional seats**: $12 / seat / month (Team), $18 / seat / month (Business)
- **Extra API requests**: $0.001 / request above plan limit
- **Premium support upgrade**: $200 / month (24h SLA on Team plan)
- **Data residency (EU)**: $500 / month (Business and above)

## Annual discount

Pay annually and get **2 months free** on Team and Business plans.

## Refund policy

Annual subscriptions can be refunded within 30 days of purchase, prorated
for the unused portion. Monthly subscriptions are non-refundable. For
Enterprise contracts, see the Master Services Agreement.
""",
        encoding="utf-8",
    )


# ── 5. faq.md ─────────────────────────────────────────────────────────────
def make_faq_md(out: Path) -> None:
    out.write_text(
        """# TechCorp FAQ

## How do I get support?

Email **support@techcorp.example** for general questions. Enterprise
customers can use the in-app chat for 1-hour SLA support.

## Where is TechCorp headquartered?

The company was founded in San Francisco and maintains its headquarters
at 500 Market Street, San Francisco, CA 94105.

## What is the maximum API rate limit?

The default rate limit is 100 requests per second per API key. Enterprise
customers can request higher limits — contact sales@techcorp.example.

## Do you have a public roadmap?

Yes, see https://roadmap.techcorp.example. We publish quarterly themes
and ship status every Friday.

## How do I install the SDK?

```bash
pip install techcorp-sdk
```

Then in your code:

```python
import techcorp
client = techcorp.Client(api_key="YOUR_KEY")
```

## What programming languages are supported?

Python, JavaScript/TypeScript, Go, and Ruby. Community-maintained SDKs
exist for Rust, Java, and C#.

## Can I export my data?

Yes. All data is exportable as JSON or CSV from the dashboard. Enterprise
customers can request full database snapshots via the support portal.

## What happens to my data if I cancel?

Account data is retained for 90 days after cancellation, then permanently
deleted. See the Terms of Service for the full retention policy.
""",
        encoding="utf-8",
    )


# ── 6. terms.html ──────────────────────────────────────────────────────────
def make_terms_html(out: Path) -> None:
    out.write_text(
        """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>TechCorp Terms of Service</title>
</head>
<body>
  <h1>TechCorp Terms of Service</h1>
  <p><em>Effective: January 1, 2025.</em></p>

  <h2>1. Acceptance of terms</h2>
  <p>By accessing or using the TechCorp platform, you agree to be bound
  by these Terms of Service. If you do not agree, you may not use the
  service.</p>

  <h2>2. Account responsibilities</h2>
  <p>You are responsible for maintaining the security of your account
  credentials and for all activity that occurs under your account. Notify
  TechCorp immediately of any unauthorized use at security@techcorp.example.</p>

  <h2>3. Data retention and deletion</h2>
  <p>Upon account termination, TechCorp will retain your data for a
  period of 90 days to allow for account recovery. After 90 days, all
  account data will be permanently deleted from production systems
  within 30 additional days. Backups are purged within 180 days total.</p>

  <h2>4. Refund policy for Enterprise customers</h2>
  <p>Enterprise customers may request a prorated refund for the unused
  portion of their annual contract within the first 90 days. Refund
  requests must be submitted to billing@techcorp.example and require
  approval by the customer's account executive.</p>

  <h2>5. Limitation of liability</h2>
  <p>To the maximum extent permitted by law, TechCorp's total liability
  for any claim arising under these terms shall not exceed the fees paid
  by you to TechCorp in the 12 months preceding the claim.</p>

  <h2>6. Governing law</h2>
  <p>These terms are governed by the laws of the State of California,
  United States, without regard to its conflict of law provisions.</p>

  <h2>7. Contact</h2>
  <p>Questions about these terms? Email legal@techcorp.example.</p>
</body>
</html>
""",
        encoding="utf-8",
    )


# ── 7. README.txt ─────────────────────────────────────────────────────────
def make_readme_txt(out: Path) -> None:
    out.write_text(
        """TechCorp Platform — README
========================

Welcome to the TechCorp developer platform. This file gives a 5-minute
overview of what the platform does and how to get started.

WHAT IS TECHCORP?
-----------------
TechCorp is a B2B SaaS platform that helps mid-market companies automate
their customer-support workflows. Our product is built around three
pillars:

  1. Ticketing — receive, route, and resolve customer issues
  2. Knowledge base — publish help articles and a public FAQ
  3. Analytics — track response time, CSAT, and agent productivity

QUICK START
-----------
1. Sign up for a free account at https://signup.techcorp.example
2. Create an API key in your dashboard under Settings > API
3. Install the SDK:  pip install techcorp-sdk
4. Make your first API call (see examples/ directory in the SDK repo)

SUPPORTED PLATFORMS
-------------------
The service runs on AWS in us-east-1, eu-west-1, and ap-southeast-1.
Enterprise customers can request a single-tenant deployment in any
AWS region.

PRICING
-------
See pricing.md for the full pricing table. We offer a free tier with
1,000 API requests per month and unlimited seats on the Free plan.

CONTACT
-------
General:    hello@techcorp.example
Support:    support@techcorp.example
Sales:      sales@techcorp.example
Security:   security@techcorp.example
Press:      press@techcorp.example
""",
        encoding="utf-8",
    )


# ── 8. QA pairs (referenced from the files above) ──────────────────────────
QA_PAIRS = [
    # (qid, question, gold_answer, source_format, source_file)
    {
        "qid": "tb_001",
        "question": "What was TechCorp's total revenue in fiscal year 2025?",
        "answer": "$48.7 million",
        "source_format": "pdf",
        "source_file": "annual_report.pdf",
    },
    {
        "qid": "tb_002",
        "question": "How many active enterprise accounts did TechCorp have at the end of 2025?",
        "answer": "1,247",
        "source_format": "pdf",
        "source_file": "annual_report.pdf",
    },
    {
        "qid": "tb_003",
        "question": "What is TechCorp's revenue target for fiscal year 2026?",
        "answer": "$62 million",
        "source_format": "pdf",
        "source_file": "annual_report.pdf",
    },
    {
        "qid": "tb_004",
        "question": "How many vacation days does a senior engineer at TechCorp get per year?",
        "answer": "25",
        "source_format": "docx",
        "source_file": "hr_handbook.docx",
    },
    {
        "qid": "tb_005",
        "question": "How many weeks of fully paid parental leave do primary caregivers get at TechCorp?",
        "answer": "16 weeks",
        "source_format": "docx",
        "source_file": "hr_handbook.docx",
    },
    {
        "qid": "tb_006",
        "question": "Which days of the week are TechCorp employees expected to be in the office?",
        "answer": "Tuesdays and Wednesdays",
        "source_format": "docx",
        "source_file": "hr_handbook.docx",
    },
    {
        "qid": "tb_007",
        "question": "Who is the Chief Technology Officer at TechCorp?",
        "answer": "Marcus Chen",
        "source_format": "csv",
        "source_file": "employees.csv",
    },
    {
        "qid": "tb_008",
        "question": "How many engineers does TechCorp have according to the 2025 annual report?",
        "answer": "184",
        "source_format": "csv",
        "source_file": "annual_report.pdf",  # cross-file: report says 184 engineers
    },
    {
        "qid": "tb_009",
        "question": "What is the monthly price of the Team plan?",
        "answer": "$49",
        "source_format": "md",
        "source_file": "pricing.md",
    },
    {
        "qid": "tb_010",
        "question": "How much does TechCorp charge per additional seat on the Business plan?",
        "answer": "$18",
        "source_format": "md",
        "source_file": "pricing.md",
    },
    {
        "qid": "tb_011",
        "question": "What is TechCorp's support email address?",
        "answer": "support@techcorp.example",
        "source_format": "md",
        "source_file": "faq.md",
    },
    {
        "qid": "tb_012",
        "question": "What is the default API rate limit?",
        "answer": "100 requests per second",
        "source_format": "md",
        "source_file": "faq.md",
    },
    {
        "qid": "tb_013",
        "question": "How many days after account cancellation does TechCorp permanently delete account data?",
        "answer": "90",
        "source_format": "html",
        "source_file": "terms.html",
    },
    {
        "qid": "tb_014",
        "question": "Within how many days of purchase can Enterprise customers request a prorated refund?",
        "answer": "90",
        "source_format": "html",
        "source_file": "terms.html",
    },
    {
        "qid": "tb_015",
        "question": "What is the TechCorp SDK installation command?",
        "answer": "pip install techcorp-sdk",
        "source_format": "txt",
        "source_file": "README.txt",
    },
    {
        "qid": "tb_016",
        "question": "What are the three pillars of the TechCorp product?",
        "answer": "Ticketing, Knowledge base, Analytics",
        "source_format": "txt",
        "source_file": "README.txt",
    },
    {
        "qid": "tb_017",
        "question": "How long is TechCorp's net income retention period mentioned in the annual report?",
        "answer": "The annual report does not mention a net income retention period; it covers fiscal year 2025 only.",
        "source_format": "pdf",
        "source_file": "annual_report.pdf",
        "is_unanswerable": True,
    },
    {
        "qid": "tb_018",
        "question": "What is TechCorp's data scientist headcount?",
        "answer": "2",
        "source_format": "csv",
        "source_file": "employees.csv",  # Ravi Krishnan + Linnea Berg
    },
    {
        "qid": "tb_019",
        "question": "How many customer support tiers does TechCorp offer?",
        "answer": "Two: community (free) and 1-hour phone (Enterprise)",
        "source_format": "md",
        "source_file": "pricing.md",
    },
    {
        "qid": "tb_020",
        "question": "What is the contract renewal period for Enterprise customers?",
        "answer": "The terms of service do not specify a contract renewal period for Enterprise customers; it would be in the Master Services Agreement.",
        "source_format": "html",
        "source_file": "terms.html",
        "is_unanswerable": True,
    },
]


def main() -> int:
    import sys
    import logging
    logging.basicConfig(level=logging.INFO, format="%(message)s")

    print(f"Generating Track B corpus into {CORPUS_DIR}")

    print("  - annual_report.pdf ...", end=" ", flush=True)
    make_annual_report_pdf(CORPUS_DIR / "annual_report.pdf")
    print("ok")

    print("  - hr_handbook.docx ...", end=" ", flush=True)
    make_hr_handbook_docx(CORPUS_DIR / "hr_handbook.docx")
    print("ok")

    print("  - employees.csv ...", end=" ", flush=True)
    make_employees_csv(CORPUS_DIR / "employees.csv")
    print("ok")

    print("  - pricing.md ...", end=" ", flush=True)
    make_pricing_md(CORPUS_DIR / "pricing.md")
    print("ok")

    print("  - faq.md ...", end=" ", flush=True)
    make_faq_md(CORPUS_DIR / "faq.md")
    print("ok")

    print("  - terms.html ...", end=" ", flush=True)
    make_terms_html(CORPUS_DIR / "terms.html")
    print("ok")

    print("  - README.txt ...", end=" ", flush=True)
    make_readme_txt(CORPUS_DIR / "README.txt")
    print("ok")

    qa_path = SCRIPT_DIR / ".cache" / "track_b_corpus" / "qa_pairs.json"
    qa_path.write_text(json.dumps(QA_PAIRS, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"  - qa_pairs.json ({len(QA_PAIRS)} pairs) ... ok")

    print(f"\nDone. Files in {CORPUS_DIR}:")
    for p in sorted(CORPUS_DIR.iterdir()):
        size = p.stat().st_size
        print(f"  {p.name:<30} {size:>8} bytes")
    return 0


if __name__ == "__main__":
    import sys
    sys.exit(main())
